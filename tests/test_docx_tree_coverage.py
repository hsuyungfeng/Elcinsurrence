"""docx 樹狀索引全語料涵蓋測試（REQ-rule-repository 驗收標準 2）。

Wave 0 預期紅燈：`docx_tree.tree_builder` 尚未實作（Plan 03 落地），
本檔案應以 ImportError/ModuleNotFoundError 收集失敗，而非語法錯誤。
"""

import glob
import os
import pytest

from config.settings import RULE_SOURCE_DIR
from elc_audit_engine.rule_repository.docx_tree import tree_builder
from elc_audit_engine.rule_repository.docx_tree.doc_converter import soffice_is_functional




requires_soffice = pytest.mark.skipif(
    not soffice_is_functional(),
    reason="soffice headless conversion unavailable (real-conversion probe failed) — .doc conversion tests need it",
)


@requires_soffice
def test_all_source_files_converted_and_processed(tmp_path):
    # NOTE: 02-01-PLAN.md's acceptance criteria and 02-RESEARCH.md both state
    # "34 total: 11 .doc + 23 .docx". Actual glob count verified this session
    # is 11 .doc + 21 .docx = 32 (not 34) -- the plan's docx figure is stale.
    # Per deviation Rule 1, this test asserts against the live glob count
    # (the real acceptance contract: "tree-build file count == source file
    # count", not a hardcoded magic number) so it doesn't spuriously fail
    # once Plan 03 lands with a correct implementation.
    doc_files = glob.glob(os.path.join(RULE_SOURCE_DIR, "*.doc"))
    docx_files = glob.glob(os.path.join(RULE_SOURCE_DIR, "*.docx"))
    total_source_files = len(doc_files) + len(docx_files)
    assert total_source_files > 0, "expected source .doc/.docx files, found none"

    staging_dir = str(tmp_path / "staging")
    trees = tree_builder.build_all_trees(RULE_SOURCE_DIR, staging_dir)
    assert len(trees) == total_source_files


@requires_soffice
def test_flat_structure_doc_produces_nested_tree(tmp_path):
    target = "2-2-7第二部第二章第七節手術-113.12.01.docx"
    staging_dir = str(tmp_path / "staging")
    trees = tree_builder.build_all_trees(RULE_SOURCE_DIR, staging_dir)

    # build_all_trees returns dict[str, dict] keyed by original source filename
    # (see 02-03-PLAN.md interfaces/key_links). Look up directly by key, with a
    # fallback scan for dict-shaped values in case keys are stored without
    # extension or in another normalized form.
    tree = trees.get(target)
    if tree is None:
        tree = next(
            (
                v
                for k, v in trees.items()
                if os.path.basename(k) == target or k == os.path.splitext(target)[0]
            ),
            None,
        )
    assert tree is not None, f"tree for {target} not found in build_all_trees output"

    def max_depth(node, current=1):
        children = getattr(node, "children", None) or node.get("children", [])
        if not children:
            return current
        return max(max_depth(child, current + 1) for child in children)

    depth = max_depth(tree)
    assert depth >= 3, f"expected at least 3 depth levels, got {depth}"


def test_table_content_merged_into_full_text(tmp_path):
    """表格內容必須併入節點 full_text（P1-3），讓表格條文可被檢索。

    若表格只在 table_refs、full_text 不含其內容，關鍵字候選／ChromaDB／
    LLM prompt 都會看不到表格承載的條文 —— 這是 46% 無匹配率的結構性
    成因。此測試直接以 python-docx 造一個「標題＋表格」文件驗證。
    """
    import docx

    docx_path = str(tmp_path / "with_table.docx")
    document = docx.Document()
    document.add_heading("第一節 檢驗", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "項目"
    table.cell(0, 1).text = "規定"
    table.cell(1, 0).text = "尿一般檢查"
    table.cell(1, 1).text = "每件給付 100 點"
    document.save(docx_path)

    from elc_audit_engine.rule_repository.docx_tree.extractor import build_tree_for_file
    tree = build_tree_for_file(docx_path, "with_table")

    assert tree["table_refs"] == []
    section = tree["children"][0]
    assert len(section["table_refs"]) == 1
    assert "尿一般檢查" in section["full_text"]
    assert "每件給付 100 點" in section["full_text"]
