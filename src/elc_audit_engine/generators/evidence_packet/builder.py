import os
import io
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from elc_audit_engine.generators.evidence_packet.image_processor import process_and_scale_image

def _fmt(value) -> str:
    """Defensive display formatter converting None/empty to em-dash placeholder."""
    if value is None or value == "":
        return "—"
    return str(value)

def _add_warning_callout(doc: Document, text: str):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run(f"⚠ {text}")
    run.font.color.rgb = RGBColor(180, 0, 0)

def build_evidence_packet_docx(
    cover_info: dict,
    tracking_data: dict,
    timeline_data: dict,
    appeal_draft: dict,
    attachment_records: list,
) -> tuple[Document, list[str]]:
    doc = Document()
    warnings = []

    # A4 Page Setup (2.54cm margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Section 1: Cover Page
    doc.add_heading("申復佐證包", level=0)
    
    doc.add_heading("Section 1: Cover Page", level=1)
    case_seq = cover_info.get("case_seq", "未知")
    doc.add_paragraph(f"案件流水號: {case_seq}")

    # Section 2: Audit Trail
    doc.add_heading("Section 2: Audit Trail", level=1)
    entries = tracking_data.get("entries", [])
    if not entries:
        doc.add_paragraph("無審核軌跡")
    for entry in entries:
        doc.add_paragraph(f"醫令: {entry.get('order_code', '')}, 狀態: {entry.get('status', '')}")

    # Section 3: Record Summary
    doc.add_heading("Section 3: Record Summary", level=1)
    events = timeline_data.get("events", [])
    if not events:
        doc.add_paragraph("無就醫紀錄")
    for event in events:
        doc.add_paragraph(f"{event.get('date', '')}: {event.get('desc', '')}")

    # Section 4: Appeal Draft
    doc.add_heading("Section 4: Appeal Draft", level=1)
    draft_sections = appeal_draft.get("sections", [])
    for sec in draft_sections:
        doc.add_heading(sec.get("title", ""), level=2)
        doc.add_paragraph(sec.get("text", ""))

    # Appendix (Images)
    if attachment_records:
        doc.add_page_break()
        doc.add_heading("Appendix: Attachments", level=1)
        
        for record in attachment_records:
            file_path = record.get("file_path")
            filename = record.get("filename", os.path.basename(file_path))
            
            doc.add_heading(f"附件: {filename}", level=2)
            
            if not file_path or not os.path.exists(file_path):
                msg = f"附件檔【{filename}】載入失敗: 檔案不存在"
                warnings.append(msg)
                _add_warning_callout(doc, msg)
                continue
                
            try:
                buf, w, h = process_and_scale_image(file_path)
                doc.add_picture(buf, width=Cm(w), height=Cm(h))
            except Exception as e:
                msg = f"附件檔【{filename}】載入失敗: {e}"
                warnings.append(msg)
                _add_warning_callout(doc, msg)

    return doc, warnings
